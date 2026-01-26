import io
import re
import zipfile
import unicodedata
import datetime

import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ============================================================
# TEMPLATE
# ============================================================
# Letak template DOCX ini dalam folder yang sama dengan app.py,
# atau upload template melalui UI (pilihan).
#
# Nama default (boleh tukar):
DEFAULT_TEMPLATE_FILENAME = "FORMAT BETUL COLLEAGUE !.docx"


# ============================================================
# Utilities
# ============================================================

_BULAN_MAP = {
    "JANUARI": 1,
    "FEBRUARI": 2,
    "MAC": 3,
    "APRIL": 4,
    "MEI": 5,
    "JUN": 6,
    "JULAI": 7,
    "OGOS": 8,
    "SEPTEMBER": 9,
    "OKTOBER": 10,
    "NOVEMBER": 11,
    "DISEMBER": 12,
}

_BULAN_TITLE = {
    1: "Januari",
    2: "Februari",
    3: "Mac",
    4: "April",
    5: "Mei",
    6: "Jun",
    7: "Julai",
    8: "Ogos",
    9: "September",
    10: "Oktober",
    11: "November",
    12: "Disember",
}

_HARI_MAP = {
    0: "Isnin",
    1: "Selasa",
    2: "Rabu",
    3: "Khamis",
    4: "Jumaat",
    5: "Sabtu",
    6: "Ahad",
}

_FIELD_PREFIXES = [
    "PEMOHON",
    "PERUNDING",
    "LOKASI",
    "KOORDINAT",
    "NO. RUJUKAN OSC",
    "NO. RUJUKAN",
    "PELAN SUSUNATUR",
    "NO FAIL",
    "NO. FAIL",
]


def _clean(s: str) -> str:
    s = s.replace("\u00a0", " ")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _to_date_string(dt: datetime.date) -> str:
    return f"{dt.day} {_BULAN_TITLE[dt.month]} {dt.year}"


def _safe_filename(s: str, max_len: int = 80) -> str:
    # Keep letters, numbers, underscore only.
    s = unicodedata.normalize("NFKD", s)
    s = re.sub(r"[\s]+", "_", s.strip())
    s = re.sub(r"[^A-Za-z0-9_]+", "", s)
    s = s.strip("_")
    if not s:
        s = "pemohon"
    return s[:max_len]


def _load_template(template_bytes: bytes | None = None) -> Document:
    """
    Load template from:
    1) uploaded bytes (if provided), else
    2) local file DEFAULT_TEMPLATE_FILENAME
    """
    if template_bytes:
        return Document(io.BytesIO(template_bytes))

    try:
        return Document(DEFAULT_TEMPLATE_FILENAME)
    except Exception as e:
        raise RuntimeError(
            f"Template '{DEFAULT_TEMPLATE_FILENAME}' tak dapat dibuka. "
            f"Sila pastikan file template ada dalam folder app.py atau upload template melalui UI. "
            f"Detail: {e}"
        )


def _set_run_font(run, name="Arial", size_pt=11, bold=None, color_rgb=(0, 0, 0)):
    """Force consistent font + BLACK color (avoid template grey)."""
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color_rgb is not None:
        run.font.color.rgb = RGBColor(*color_rgb)


def _set_paragraph_text(paragraph, text: str, name="Arial", size_pt=11, bold=False, align=None):
    # Clear existing runs but keep paragraph formatting.
    for r in list(paragraph.runs):
        r.text = ""
    paragraph.text = ""
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    _set_run_font(run, name=name, size_pt=size_pt, bold=bold)
    return paragraph


def _set_cell_text(
    cell,
    text: str,
    name="Arial",
    size_pt=11,
    bold=False,
    preserve_linebreaks=True,
    align=WD_ALIGN_PARAGRAPH.LEFT,
):
    # Clear existing paragraphs
    for p in cell.paragraphs:
        for r in list(p.runs):
            r.text = ""
        p.text = ""
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Use first paragraph
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    if preserve_linebreaks and "\n" in text:
        run = p.add_run("")
        _set_run_font(run, name=name, size_pt=size_pt, bold=bold)
        parts = text.split("\n")
        for idx, part in enumerate(parts):
            if idx > 0:
                run.add_break()
            run.add_text(part)
    else:
        run = p.add_run(text)
        _set_run_font(run, name=name, size_pt=size_pt, bold=bold)


def _parse_meeting_info(paras: list[str]) -> tuple[str | None, datetime.date | None]:
    full = "\n".join(paras)

    m_bil = re.search(r"\bBIL\.?\s*[:.-]?\s*(\d{1,2}/\d{4})", full, flags=re.IGNORECASE)
    bil = m_bil.group(1) if m_bil else None

    # Find first date in Malay uppercase form (e.g., 12 JANUARI 2026)
    m_date = re.search(r"(\d{1,2})\s+([A-Z]{3,})\s+(\d{4})", full)
    if not m_date:
        return bil, None

    day = int(m_date.group(1))
    mon_name = m_date.group(2).upper()
    year = int(m_date.group(3))
    if mon_name not in _BULAN_MAP:
        return bil, None

    try:
        dt = datetime.date(year, _BULAN_MAP[mon_name], day)
    except Exception:
        dt = None
    return bil, dt


def _split_blocks(paras: list[str]) -> list[tuple[str, list[str]]]:
    blocks = []
    i = 0
    while i < len(paras):
        line = paras[i]
        if line.upper().startswith("KERTAS MESYUARAT BIL."):
            kertas = re.sub(r"^KERTAS MESYUARAT BIL\.?\s*", "", line, flags=re.IGNORECASE).strip()
            j = i + 1
            block_lines = []
            while j < len(paras) and not paras[j].upper().startswith("KERTAS MESYUARAT BIL."):
                block_lines.append(paras[j])
                j += 1
            blocks.append((kertas, block_lines))
            i = j
        else:
            i += 1
    return blocks


def _extract_value(line: str) -> str:
    # supports "Pemohon : X" or "Pemohon\t:\tX"
    parts = re.split(r"\s*:\s*", line, maxsplit=1)
    if len(parts) == 2:
        return parts[1].strip()
    return re.sub(r"^[A-Za-z .()/-]+\s+", "", line).strip()


def _strip_parentheses(s: str) -> str:
    return re.sub(r"\s*\([^)]*\)", "", s).strip()


def _parse_case(kertas: str, block_lines: list[str]) -> dict | None:
    kertas_clean = _clean(kertas)

    if "OSC/PKM/" in kertas_clean:
        jenis = "Kebenaran Merancang"
        jenis_code = "PKM"
    elif "OSC/BGN/" in kertas_clean:
        jenis = "Bangunan"
        jenis_code = "BGN"
    else:
        return None

    # Description/Nama Permohonan = lines from start until first field-prefix line
    desc_lines = []
    for raw in block_lines:
        l = raw.strip()
        if not l:
            continue
        if l.upper().startswith(tuple(_FIELD_PREFIXES)):
            break
        desc_lines.append(raw.rstrip())

    nama_permohonan = "\n".join([dl.strip() for dl in desc_lines if dl.strip()]).strip()

    pemohon = ""
    perunding = ""
    no_ruj_osc = ""

    for raw in block_lines:
        l = raw.strip()
        if not l:
            continue
        u = l.upper()
        if u.startswith("PEMOHON"):
            pemohon = _strip_parentheses(_extract_value(l))
        elif u.startswith("PERUNDING"):
            perunding_raw = _strip_parentheses(_extract_value(l))
            perunding = re.split(r"\s+Ar\.", perunding_raw, maxsplit=1)[0].strip() or perunding_raw
        elif u.startswith("NO. RUJUKAN OSC"):
            no_ruj_osc = _extract_value(l).strip()

    return {
        "kertas": kertas_clean,
        "jenis_code": jenis_code,
        "jenis_permohonan": jenis,
        "pemohon": pemohon,
        "perunding": perunding,
        "nama_permohonan": nama_permohonan,
        "id_permohonan": no_ruj_osc,
    }


def parse_agenda(doc: Document) -> dict:
    paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    paras = [p.replace("\u00a0", " ").rstrip() for p in paras]

    meeting_bil, meeting_date = _parse_meeting_info([_clean(p) for p in paras])

    blocks = _split_blocks(paras)
    cases = []
    for kertas, lines in blocks:
        c = _parse_case(kertas, lines)
        if c:
            cases.append(c)

    return {
        "meeting_bil": meeting_bil,
        "meeting_date": meeting_date,
        "cases": cases,
    }


# ============================================================
# Formatting fixes (based on your screenshots)
# ============================================================

def _remove_leading_blank_paragraphs(doc: Document, max_remove: int = 10):
    """
    Buang perenggan kosong di paling atas body supaya lepas header tak jadi ada "single ruang".
    """
    removed = 0
    while removed < max_remove and doc.paragraphs:
        p = doc.paragraphs[0]
        if p.text.strip():
            break
        p._element.getparent().remove(p._element)
        removed += 1


def _ensure_header_lines(doc: Document, rujukan_kami: str, tarikh_str: str):
    """
    Pastikan info header wujud dalam Word Header (jadi auto repeat page 2+)
    dan paksa warna hitam.
    """
    for section in doc.sections:
        hdr = section.header

        # Update existing lines if present
        found_ruj_tuan = False
        found_ruj_kami = False
        found_tarikh = False

        for p in hdr.paragraphs:
            t = (p.text or "").strip()
            if t.startswith("Rujukan Tuan"):
                _set_paragraph_text(p, "Rujukan Tuan :", align=WD_ALIGN_PARAGRAPH.RIGHT)
                found_ruj_tuan = True
            elif t.startswith("Rujukan Kami"):
                _set_paragraph_text(p, f"Rujukan Kami : {rujukan_kami}", align=WD_ALIGN_PARAGRAPH.RIGHT)
                found_ruj_kami = True
            elif t.startswith("Tarikh"):
                _set_paragraph_text(p, f"Tarikh : {tarikh_str}", align=WD_ALIGN_PARAGRAPH.RIGHT)
                found_tarikh = True

        # If not found, create
        if not found_ruj_tuan:
            p = hdr.add_paragraph()
            _set_paragraph_text(p, "Rujukan Tuan :", align=WD_ALIGN_PARAGRAPH.RIGHT)
        if not found_ruj_kami:
            p = hdr.add_paragraph()
            _set_paragraph_text(p, f"Rujukan Kami : {rujukan_kami}", align=WD_ALIGN_PARAGRAPH.RIGHT)
        if not found_tarikh:
            p = hdr.add_paragraph()
            _set_paragraph_text(p, f"Tarikh : {tarikh_str}", align=WD_ALIGN_PARAGRAPH.RIGHT)

        # Reduce header spacing
        for p in hdr.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)


def _fix_checkbox_alignment(doc: Document):
    """
    Selarikan kotak (LULUS/LULUS BERSYARAT) dengan (TOLAK/TANGGUH),
    bila struktur template adalah table.
    """
    target = None
    for tbl in doc.tables:
        txt = "\n".join([c.text for r in tbl.rows for c in r.cells]).upper()
        if "LULUS" in txt and "TOLAK" in txt and "TANGGUH" in txt:
            target = tbl
            break

    if not target:
        return

    for row in target.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf = p.paragraph_format
                pf.left_indent = Cm(0)
                pf.first_line_indent = Cm(0)
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)

                # buang leading spaces/tabs yang buat "lari"
                if p.text and p.text[:1].isspace():
                    _set_paragraph_text(p, p.text.lstrip(), align=WD_ALIGN_PARAGRAPH.LEFT)


# ============================================================
# Document generation
# ============================================================

def _fill_doc(
    case: dict,
    meeting_bil: str,
    meeting_date: datetime.date,
    case_index: int,
    template_bytes: bytes | None,
) -> Document:
    doc = _load_template(template_bytes)

    tarikh_str = _to_date_string(meeting_date)
    hari_str = _HARI_MAP[meeting_date.weekday()]

    # IMPORTANT: Rujukan Kami ikut urutan jana dokumen (1..N), bukan ikut "001/2026" dsb
    year = meeting_bil.split("/")[1] if meeting_bil and "/" in meeting_bil else str(meeting_date.year)
    rujukan_kami = f"({case_index})MBSP/15/1551/(   ){year}"

    # Ensure header repeats on page 2+ and text black
    _ensure_header_lines(doc, rujukan_kami=rujukan_kami, tarikh_str=tarikh_str)

    # Remove extra space below header (body top)
    _remove_leading_blank_paragraphs(doc)

    # If template stores these lines in body (not header), fix them too (black + right align)
    for p in doc.paragraphs[:25]:
        t = (p.text or "").strip()
        if t.startswith("Rujukan Tuan"):
            _set_paragraph_text(p, "Rujukan Tuan :", align=WD_ALIGN_PARAGRAPH.RIGHT)
            p.paragraph_format.space_after = Pt(0)
        elif t.startswith("Rujukan Kami"):
            _set_paragraph_text(p, f"Rujukan Kami : {rujukan_kami}", align=WD_ALIGN_PARAGRAPH.RIGHT)
            p.paragraph_format.space_after = Pt(0)
        elif t.startswith("Tarikh"):
            _set_paragraph_text(p, f"Tarikh : {tarikh_str}", align=WD_ALIGN_PARAGRAPH.RIGHT)
            p.paragraph_format.space_after = Pt(0)

    # Top table (Kepada/Pemilik/Jenis/Nama/ID)
    if doc.tables:
        t0 = doc.tables[0]
        _set_cell_text(t0.cell(0, 2), case.get("perunding", ""), preserve_linebreaks=False, align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_text(t0.cell(1, 2), case.get("pemohon", ""), preserve_linebreaks=False, align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_text(t0.cell(2, 2), case.get("jenis_permohonan", ""), preserve_linebreaks=False, align=WD_ALIGN_PARAGRAPH.LEFT)

        # IMPORTANT: Nama Permohonan MUST be LEFT aligned (bukan center)
        _set_cell_text(
            t0.cell(3, 2),
            case.get("nama_permohonan", ""),
            preserve_linebreaks=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )

        _set_cell_text(t0.cell(4, 2), case.get("id_permohonan", ""), preserve_linebreaks=False, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Paragraph 2 (Mesyuarat bil & tarikh) — rebuild runs (bold) and BLACK
    target_para = None
    for p in doc.paragraphs:
        if "Adalah dimaklumkan bahawa" in p.text and "Mesyuarat Jawatankuasa Pusat Setempat" in p.text:
            target_para = p
            break

    if target_para is not None:
        for r in list(target_para.runs):
            r.text = ""
        target_para.text = ""

        def add(text, bold=False):
            r = target_para.add_run(text)
            _set_run_font(r, name="Arial", size_pt=11, bold=bold, color_rgb=(0, 0, 0))
            return r

        add("2.\tAdalah dimaklumkan bahawa ")
        add("Mesyuarat Jawatankuasa Pusat Setempat (OSC) ", bold=True)
        add(f"Bil.{meeting_bil} ", bold=True)
        add("yang bersidang pada ")
        add(f"{tarikh_str} ({hari_str}) ", bold=True)
        add(
            "bersetuju untuk memberikan keputusan ke atas permohonan yang telah dikemukakan oleh pihak tuan/ puan seperti mana berikut:"
        )

    # Checkbox alignment fix
    _fix_checkbox_alignment(doc)

    return doc


def build_zip(agenda_doc: Document, template_bytes: bytes | None) -> tuple[bytes, dict]:
    parsed = parse_agenda(agenda_doc)
    meeting_bil = parsed["meeting_bil"]
    meeting_date = parsed["meeting_date"]
    cases = parsed["cases"]

    if not meeting_bil or not meeting_date:
        raise ValueError("Tak jumpa BIL. (cth: 01/2026) atau tarikh mesyuarat (cth: 12 JANUARI 2026) dalam agenda.")

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for idx, case in enumerate(cases, start=1):
            doc = _fill_doc(case, meeting_bil, meeting_date, case_index=idx, template_bytes=template_bytes)

            pemohon_short = _safe_filename(case.get("pemohon", "pemohon"))
            k = case.get("kertas", "").replace("/", "_")
            filename = f"{idx:02d}_{k}_{pemohon_short}.docx"

            fbuf = io.BytesIO()
            doc.save(fbuf)
            z.writestr(filename, fbuf.getvalue())

    meta = {
        "meeting_bil": meeting_bil,
        "meeting_date": _to_date_string(meeting_date),
        "count_total": len(cases),
        "count_pkm": sum(1 for c in cases if c.get("jenis_code") == "PKM"),
        "count_bgn": sum(1 for c in cases if c.get("jenis_code") == "BGN"),
    }
    return buf.getvalue(), meta


# ============================================================
# Streamlit UI
# ============================================================

st.set_page_config(page_title="Pemakluman Keputusan JKOSC", layout="centered")
st.title("Pemakluman Keputusan JKOSC")

template_upload = st.file_uploader("Optional: Upload Template (DOCX)", type=["docx"])
agenda_upload = st.file_uploader("Upload fail Agenda (DOCX)", type=["docx"])

template_bytes = None
if template_upload is not None:
    template_bytes = template_upload.read()

if agenda_upload:
    try:
        agenda_doc = Document(io.BytesIO(agenda_upload.read()))
        parsed = parse_agenda(agenda_doc)

        meeting_bil = parsed["meeting_bil"]
        meeting_date = parsed["meeting_date"]
        cases = parsed["cases"]

        if meeting_bil:
            st.success(f"Jumpa BIL. {meeting_bil}")
        else:
            st.warning("Tak jumpa BIL. (contoh: 01/2026).")

        if meeting_date:
            st.success(f"Jumpa tarikh mesyuarat: {_to_date_string(meeting_date)} ({_HARI_MAP[meeting_date.weekday()]})")
        else:
            st.warning("Tak jumpa tarikh mesyuarat (contoh: 12 JANUARI 2026).")

        st.info(
            f"Kes PKM dijumpai: {sum(1 for c in cases if c.get('jenis_code')=='PKM')} | "
            f"Kes BGN dijumpai: {sum(1 for c in cases if c.get('jenis_code')=='BGN')} | "
            f"Jumlah: {len(cases)}"
        )

        if st.button("Download semua dokumen (ZIP)"):
            zip_bytes, meta = build_zip(agenda_doc, template_bytes=template_bytes)
            zip_name = f"pemakluman_keputusan_{meeting_bil.replace('/','_')}.zip"

            st.download_button(
                label=f"Klik untuk download ZIP ({meta['count_total']} dokumen)",
                data=zip_bytes,
                file_name=zip_name,
                mime="application/zip",
            )

    except Exception as e:
        st.error(str(e))
